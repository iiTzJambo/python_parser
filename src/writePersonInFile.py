from docx import Document
from docx.shared import Inches


personsPath = "persons/"


def saveInFile(person, fileName):
    if (person["name"] is None):
        return

    doc = Document()

    print(
        f'Person name: {person['name']}', '\n'
    )

    doc.add_heading(person["name"], 1)

    doc.add_picture(person['photo'], width=Inches(2.7), height=Inches(3))

    doc.add_paragraph(person["personalInfo"])

    doc.save(f"{personsPath}/{fileName}.docx")
