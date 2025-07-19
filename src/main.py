import requests 
from bs4 import BeautifulSoup
from fake_useragent import UserAgent 
from translate import Translator
from docx import Document

ua = UserAgent() 
headers = {"User-Agent": ua.random} #Этот объект позволяет генерировать случайные User-Agent строки, которые имитируют реальные браузеры (Chrome, Firefox, Safari и др.).

translator = Translator(to_lang="ru")#Строка для дальнейшего перевода собранной инфы
 
response_1 = requests.get("https://www.jbcharleston.jb.mil/About-Us/Biographies/Display/Article/3094628/colonel-jason-h-jp-parker/", headers=headers)#запрос на сайт с биографией
response_2 = requests.get("https://www.jbcharleston.jb.mil/About-Us/Biographies/Display/Article/4233762/chief-master-sergerant-richard-b-ricky-lamm/", headers=headers)
response_3 = requests.get("https://www.jbcharleston.jb.mil/About-Us/Biographies/Display/Article/2656971/captain-g-reed-koepp-ii/",headers=headers)
response_4 = requests.get("https://www.jbcharleston.jb.mil/About-Us/Biographies/Display/Article/1756649/command-master-chief-joel-brandt/",headers=headers)
response_5 = requests.get("https://www.jbcharleston.jb.mil/About-Us/Biographies/Display/Article/3209010/mr-ray-a-forcier/",headers=headers)

soup_1 = BeautifulSoup(response_1.text, "lxml")#строка для парсинга первого сайта, собирает инфу из HTML
soup_2 = BeautifulSoup(response_2.text, "lxml")
soup_3 = BeautifulSoup(response_3.text, "lxml")
soup_4 = BeautifulSoup(response_4.text, "lxml")
soup_5 = BeautifulSoup(response_5.text, "lxml") 

def infa_name(soup_parametr):#функция которая возвращает "имя" военного из HTML кода  
    main_div = soup_parametr.find("div",class_ = "border-wrapper")
    soldier_name_h = soup_parametr.find("h1", class_ = "title").text
    return soldier_name_h

def infa_img (soup_parametr):#функция возвращает ссылку на картинку
    main_div = soup_parametr.find("div",class_ = "border-wrapper")
    img_src = main_div.find("img",class_="img-responsive").get("src")
    return img_src

def infa_org_name(soup_parametr):#фукнция будет возвращать название организации в которой служит военный
    main_div = soup_parametr.find("div",class_ = "border-wrapper")
    soldier_biography = main_div.find("div", class_="body").text
    return soldier_biography




#Строки кода ниже изменяют (текст) в файл шаблона для заполнения и переименовывают от первого до пятого.
#Док. 1 Док. 2 и т.д
doc = Document("Шаблон для заполнения — копия.docx")
i=1
def writer_1(name,photo,org_name,job,rank,exp,educ,bio,source,create):#функция заменяет (текст) в пунктах шаблона и после создаёт новый документ и изменёным текстом
    doc.paragraphs[1].text = name
    doc.paragraphs[2].text = image
    doc.paragraphs[4].text = org_name
    doc.paragraphs[6].text = work
    doc.paragraphs[8].text = job
    doc.paragraphs[10].text = rank
    doc.paragraphs[12].text = exp
    doc.paragraphs[14].text = educ
    doc.paragraphs[16].text = bio
    doc.paragraphs[18].text = source
    doc.paragraphs[20].text = day_doc
    doc.save(f"Шаблон для заполнения_{i}.docx")
def writer_2(name,photo,org_name,work,job,rank,exp,educ,bio,source,day_doc):
    doc.paragraphs[1].text = name
    doc.paragraphs[2].text = image
    doc.paragraphs[4].text = org_name
    doc.paragraphs[6].text = work
    doc.paragraphs[8].text = job
    doc.paragraphs[10].text = rank
    doc.paragraphs[12].text = exp
    doc.paragraphs[14].text = educ
    doc.paragraphs[16].text = bio
    doc.paragraphs[18].text = source
    doc.paragraphs[20].text = day_doc
    doc.save(f"Шаблон для заполнения_{i+1}.docx")
def writer_3(name,photo,org_name,work,job,rank,exp,educ,bio,source,day_doc):
    doc.paragraphs[1].text = name
    doc.paragraphs[2].text = image
    doc.paragraphs[4].text = org_name
    doc.paragraphs[6].text = work
    doc.paragraphs[8].text = job
    doc.paragraphs[10].text = rank
    doc.paragraphs[12].text = exp
    doc.paragraphs[14].text = educ
    doc.paragraphs[16].text = bio
    doc.paragraphs[18].text = source
    doc.paragraphs[20].text = day_doc
    doc.save(f"Шаблон для заполнения_{i+2}.docx")
def writer_4(name,photo,org_name,work,job,rank,exp,educ,bio,source,day_doc):
    doc.paragraphs[1].text = name
    doc.paragraphs[2].text = image
    doc.paragraphs[4].text = org_name
    doc.paragraphs[6].text = work
    doc.paragraphs[8].text = job
    doc.paragraphs[10].text = rank
    doc.paragraphs[12].text = exp
    doc.paragraphs[14].text = educ
    doc.paragraphs[16].text = bio
    doc.paragraphs[18].text = source
    doc.paragraphs[20].text = day_doc
    doc.save(f"Шаблон для заполнения_{i+3}.docx")       
def writer_5(name,photo,org_name,work,job,rank,exp,educ,bio,source,day_doc):
    doc.paragraphs[1].text = name
    doc.paragraphs[2].text = image
    doc.paragraphs[4].text = org_name
    doc.paragraphs[6].text = work
    doc.paragraphs[8].text = job
    doc.paragraphs[10].text = rank
    doc.paragraphs[12].text = exp
    doc.paragraphs[14].text = educ
    doc.paragraphs[16].text = bio
    doc.paragraphs[18].text = source
    doc.paragraphs[20].text = day_doc
    doc.save(f"Шаблон для заполнения_{i+4}.docx")

writer_1("name","photo","org_name","work","job","rank","exp","educ","bio","source","day_doc")
writer_2("name","photo","org_name","work","job","rank","exp","educ","bio","source","day_doc")
writer_3("name","photo","org_name","work","job","rank","exp","educ","bio","source","day_doc")
writer_4("name","photo","org_name","work","job","rank","exp","educ","bio","source","day_doc")
writer_5("name","photo","org_name","work","job","rank","exp","educ","bio","source","day_doc")